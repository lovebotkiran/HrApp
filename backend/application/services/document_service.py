import os
try:
    import docx
except ImportError:
    docx = None
from datetime import datetime
from typing import Dict, Any
import uuid

class DocumentService:
    def __init__(self, template_dir: str = "resources/templates"):
        # Get the absolute path to the templates directory
        # Assuming this file is in backend/application/services/
        base_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        self.template_dir = os.path.join(base_dir, template_dir)
        self.output_dir = os.path.join(base_dir, "uploads", "documents")
        
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir, exist_ok=True)

    def generate_document(self, template_name: str, data: Dict[str, Any], output_filename: str = None) -> str:
        """
        Generates a document from a template by replacing placeholders.
        Returns the path to the generated document.
        """
        template_path = os.path.join(self.template_dir, template_name)
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Template not found: {template_path}")

        if not docx:
            raise ImportError("python-docx is not installed unless offer generation is enabled")

        doc = docx.Document(template_path)
        
        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            self._replace_placeholders(para, data)
            
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_placeholders(para, data)

        if not output_filename:
            output_filename = f"{uuid.uuid4()}.docx"
        
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        
        return output_path

    def _replace_placeholders(self, paragraph, data: Dict[str, Any]):
        """
        Replace {{KEY}} with value in paragraph text.
        """
        for key, value in data.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in paragraph.text:
                # To preserve formatting as much as possible, we try to preserve the runs
                # But simple replacement in text is often safer for template tags
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    def get_offer_letter_data(self, offer, candidate, job_posting) -> Dict[str, Any]:
        """
        Prepare data for offer letter template.
        """
        return {
            "OFFER_DATE": datetime.now().strftime("%d %B %Y"),
            "CANDIDATE_NAME": f"{candidate.first_name} {candidate.last_name}",
            "CANDIDATE_ADDRESS": candidate.current_location or "Address Not Provided",
            "POSITION": offer.designation or job_posting.title,
            "START_DATE": offer.joining_date.strftime("%d %B %Y") if offer.joining_date else "TBD",
            "COMPENSATION": f"{offer.annual_ctc:,.2f}" if offer.annual_ctc else "0.00",
            "EMPLOYMENT_TYPE": offer.employment_type or "Full-time",
            "OFFER_DATE_DEADLINE": offer.offer_valid_until.strftime("%d %B %Y") if offer.offer_valid_until else "TBD"
        }

    def get_nda_data(self, offer, candidate, job_posting) -> Dict[str, Any]:
        """
        Prepare data for NDA template.
        """
        return {
            "AGREEMENT_DATE": datetime.now().strftime("%d %B %Y"),
            "JURISDICTION": "Sri Lanka", # Should ideally be configurable
            "CANDIDATE_NAME": f"{candidate.first_name} {candidate.last_name}",
            "CANDIDATE_ADDRESS": candidate.current_location or "Address Not Provided",
            "POSITION": offer.designation or job_posting.title,
        }
